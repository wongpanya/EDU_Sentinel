from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("EDU_Sentinel_System_Development_Work_Order_v1.1.docx")

FONT = "Arial"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 31, 43)
MUTED = RGBColor(90, 100, 110)
LIGHT = "E8EEF5"
GRAY = "F2F4F7"
PALE = "F8FAFC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_para(doc, text="", style=None, size=11, bold=False, color=INK, after=6, before=0):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13, color=BLUE if level < 3 else DARK_BLUE, bold=True)
    return p


def add_table(doc, headers, rows, widths, header_fill=LIGHT):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], header_fill)
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run_font(r, size=9.5, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(val)
            set_run_font(r, size=9.5, color=INK)
            set_cell_margins(cells[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def add_masthead(doc):
    header = doc.sections[0].header.paragraphs[0]
    header.text = "EDU Sentinel | Development Work Order"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        set_run_font(run, size=9, color=MUTED)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "Prepared for prototype development"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, size=9, color=MUTED)

    add_para(doc, "DEVELOPMENT WORK ORDER", size=10, bold=True, color=MUTED, after=2)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("EDU Sentinel")
    set_run_font(r, size=28, color=INK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run("System Development Blueprint for Starting Prototype Implementation")
    set_run_font(r, size=14, color=MUTED)

    add_table(
        doc,
        ["รายการ", "รายละเอียด"],
        [
            ("วัตถุประสงค์", "ใช้เป็นคำสั่งงานตั้งต้นสำหรับทีม Product, UX/UI, Data, Backend, Frontend และ AI เพื่อเริ่มพัฒนา Prototype"),
            ("ขอบเขตเริ่มต้น", "Phase 1 Executive MVP: dashboard, map, alert, data import, AI summary และ report export"),
            ("เอกสารอ้างอิง", "EDU Sentinel System Blueprint v1.0 และ Development Blueprint Summary"),
            ("สถานะ", "พร้อมใช้สำหรับ kickoff, backlog breakdown และ sprint planning"),
        ],
        [1.55, 4.95],
        header_fill=GRAY,
    )


def main():
    doc = Document()
    style_doc(doc)
    add_masthead(doc)

    add_heading(doc, "1. Product Intent", 1)
    add_para(
        doc,
        "EDU Sentinel ต้องถูกพัฒนาเป็นระบบ Policy Intelligence + Early Warning ที่เปลี่ยนข้อมูลด้านการศึกษา สังคม เศรษฐกิจ พื้นที่ และมาตรการช่วยเหลือ ให้เป็นสัญญาณที่ผู้บริหารเห็นเร็ว เข้าใจสาเหตุ และสั่งการต่อได้ทันที",
    )
    add_bullet(doc, "ไม่ใช่ dashboard แสดงรายงานย้อนหลัง แต่เป็นวงจร Data > Intelligence > Alert > Decision > Action > Learning")
    add_bullet(doc, "AI เป็น Analyst Copilot ที่อธิบายพร้อมหลักฐาน ไม่ใช่ Automated Decision Maker")
    add_bullet(doc, "ทุก KPI และ Alert ต้องตรวจสอบ source, definition, period, data quality และ rule/model version ได้")

    add_para(
        doc,
        "Product operating model: ทุกหน้าจอและ API ต้องช่วยให้ผู้บริหารเห็นสัญญาณเร็ว เข้าใจสาเหตุ ตัดสินใจเชิงนโยบายได้ และส่งต่อเป็น action ที่ติดตามผลได้ โดยระบบต้องเก็บหลักฐานและ metadata ให้ตรวจสอบย้อนหลังได้เสมอ",
        bold=True,
        color=DARK_BLUE,
        after=8,
    )
    add_table(
        doc,
        ["Layer", "Requirement", "Traceability Check"],
        [
            ("Data", "รวบรวมข้อมูลการศึกษา สังคม เศรษฐกิจ พื้นที่ และมาตรการช่วยเหลือ พร้อม lineage", "source_id, period, freshness, completeness, owner"),
            ("Intelligence", "แปลงข้อมูลเป็น KPI, risk driver, coverage gap และ what-changed insight", "definition, formula, denominator, aggregation level"),
            ("Alert", "แจ้งเตือนจาก rule/model ที่ versioned และอธิบายเหตุผลได้", "rule_version/model_version, trigger, severity, confidence"),
            ("Decision", "นำเสนอ evidence, limitation และ recommended next action สำหรับผู้บริหาร", "evidence_refs, data_quality_flag, analyst note"),
            ("Action", "เปิด case/action จาก alert พร้อม owner, SLA, status และ outcome", "case_id, owner, due_at, status, resolution"),
            ("Learning", "บันทึกผลลัพธ์เพื่อปรับ rule, threshold, workflow และ prompt/evaluation", "outcome, feedback, review_date, change_log"),
        ],
        [1.05, 3.35, 2.1],
        header_fill=GRAY,
    )

    add_heading(doc, "2. Phase 1 Scope: Executive MVP", 1)
    add_table(
        doc,
        ["ลำดับ", "Feature", "เป้าหมาย", "Acceptance Criteria"],
        [
            ("P1", "Authentication & Roles", "เข้าใช้งานตามบทบาท", "มี mock/real login, role Executive/Analyst/Admin, route protection"),
            ("P2", "Data Import", "นำเข้าข้อมูลตัวอย่าง", "รองรับ CSV/Excel, validate schema, แสดง error/import status และบันทึก source/period/data quality"),
            ("P3", "Executive Command Center", "เห็นสถานการณ์ใน 1 หน้าจอ", "มี KPI, trend, top alerts, what changed, coverage gap, filter เวลา/พื้นที่ และ metadata ตรวจสอบนิยาม/แหล่งข้อมูล"),
            ("P4", "Thailand Risk Map", "เห็นพื้นที่เสี่ยง", "แสดง heatmap ระดับจังหวัด/อำเภอ/โรงเรียน และ drill-down ได้"),
            ("P5", "Early Warning Engine", "สร้าง alert จาก rule", "มี rule-based threshold/trend, severity, confidence, explanation และ rule/model version"),
            ("P6", "Alert Center & Detail", "จัดลำดับและตรวจหลักฐาน", "list/filter/sort, detail แสดง drivers, data source, period, data quality, status, SLA"),
            ("P7", "AI Executive Summary", "สรุปสถานการณ์จากข้อมูลจริง", "ตอบเป็น Answer/Evidence/Why/Limitations/Next Action พร้อม citation และ model/prompt version"),
            ("P8", "Executive Report Export", "ส่งออก brief", "export PDF/DOCX หรือ HTML print view จาก dashboard/alert"),
        ],
        [0.55, 1.45, 1.85, 2.65],
    )

    add_heading(doc, "3. Recommended Team Backlog", 1)
    add_table(
        doc,
        ["ทีม", "งานเริ่มต้น", "ผลลัพธ์ที่ต้องส่ง"],
        [
            ("Product/BA", "นิยาม persona, user journey, KPI dictionary, alert definition, demo storyline", "PRD v0.1, KPI/Alert catalog, prioritized backlog"),
            ("UX/UI", "ออกแบบ Command Center, Map, Alert Detail, AI Analyst, Case Tracking", "Figma/wireframe, component states, responsive desktop spec"),
            ("Data", "กำหนด sample dataset, schema, mapping, data quality rules", "Data dictionary, seed dataset, validation rules"),
            ("Backend", "สร้าง API domain: dashboard, areas, alerts, cases, ai, data-health, admin", "OpenAPI spec, service skeleton, DB migration"),
            ("Frontend", "ตั้งค่า app shell, routing, charts, map, table, filters", "Clickable UI เชื่อม mock API และ state จริง"),
            ("AI", "ออกแบบ prompt contract, retrieval boundary, citation format, guardrails", "AI summary endpoint, evaluation cases, refusal/limitation patterns"),
            ("QA/Security", "กำหนด test cases, RBAC tests, audit log checks, data masking checks", "Test plan, regression checklist, security checklist"),
        ],
        [1.05, 3.05, 2.4],
    )

    add_heading(doc, "4. Sprint 0: ก่อนเขียนฟีเจอร์", 1)
    for item in [
        "ตั้ง repository structure, environment, CI, lint/test, branch strategy และ issue template",
        "ยืนยัน stack: Next.js/React, API service, PostgreSQL/PostGIS หรือ mock DB สำหรับ prototype",
        "สร้าง design system ขั้นต้น: layout, navigation, card/table/chart/map states, severity colors",
        "สร้าง sample dataset สำหรับ province/district/school/student aggregate/alert/case/policy coverage",
        "ล็อก KPI dictionary และ alert rules ชุดแรก เช่น absenteeism, dropout risk, coverage gap, SLA breach",
        "นิยาม RBAC matrix และ data masking rule สำหรับ executive/analyst/case worker/admin",
        "เขียน OpenAPI draft และ mock response สำหรับทุกหน้าจอ Phase 1",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5. Minimum Data Contract", 1)
    add_table(
        doc,
        ["Entity", "Required Fields", "ใช้ในหน้าจอ/งาน"],
        [
            ("Area", "area_id, name, level, parent_id, geo_boundary, risk_index", "Map, dashboard filter, drill-down"),
            ("School", "school_id, name, type, area_id, student_count, location", "School profile, map point, alert context"),
            ("Indicator", "indicator_id, area_id/school_id, period, value, denominator, source_id", "KPI, trend, evidence"),
            ("RiskAssessment", "target_type, target_id, score, level, drivers, rule_version, confidence", "Risk pulse, map, alert explanation"),
            ("Alert", "alert_id, type, severity, status, priority_score, trigger, evidence_refs, sla_due_at", "Alert center/detail, executive top alerts"),
            ("CaseAction", "case_id, alert_id, owner, status, action_note, due_at, resolved_at, outcome", "Action tracking, SLA, learning loop"),
            ("PolicyProgram", "program_id, target_group, eligibility_rule, area_scope, coverage_count", "Coverage gap, policy intelligence"),
            ("DataSource", "source_id, name, refreshed_at, completeness, freshness, lineage", "Data health and trust indicators"),
        ],
        [1.35, 3.05, 2.1],
    )

    add_heading(doc, "6. API Domains to Start", 1)
    add_table(
        doc,
        ["Domain", "Endpoint ตัวอย่าง", "หมายเหตุ"],
        [
            ("/dashboard", "GET /summary, GET /risk-pulse, GET /what-changed", "ใช้ aggregate data และ filter เวลา/พื้นที่"),
            ("/areas", "GET /areas, GET /areas/{id}/profile, GET /areas/{id}/children", "รองรับ map drill-down"),
            ("/alerts", "GET /alerts, GET /alerts/{id}, POST /alerts/{id}/ack", "ต้องคืน evidence, drivers, status, SLA"),
            ("/cases", "POST /cases, PATCH /cases/{id}, GET /cases?alert_id=", "เชื่อม alert เป็น action"),
            ("/policies", "GET /programs, GET /coverage-gap", "Phase 1 ใช้ mock/seed ได้ แต่ schema ต้องรองรับต่อ"),
            ("/ai", "POST /executive-summary, POST /explain-alert", "ต้อง permission-aware และ cite evidence"),
            ("/data-health", "GET /sources, GET /quality", "แสดง completeness/freshness/lineage"),
        ],
        [1.1, 2.6, 2.8],
    )

    add_heading(doc, "7. UI Screens and User Flow", 1)
    add_para(doc, "Demo flow ที่ต้องทำให้เดินได้ใน 5-7 นาที:")
    for item in [
        "Executive เปิด Command Center แล้วเห็น Critical Alerts, Risk Pulse และพื้นที่ที่แย่ลง",
        "คลิก hotspot บนแผนที่เพื่อ drill-down จากประเทศ > จังหวัด > อำเภอ > โรงเรียน",
        "เปิด Alert Detail เพื่อดู what happened, affected population, drivers, trend, evidence และ data quality",
        "กดถาม AI ว่าทำไมพื้นที่นี้น่ากังวล และมีมาตรการใดครอบคลุมหรือยังตกหล่น",
        "สร้าง Case/Action จาก Alert, assign owner, ตั้ง SLA และบันทึก next action",
        "กลับ Command Center แล้วเห็น action status/SLA เปลี่ยนตาม workflow",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "8. AI Guardrails", 1)
    add_table(
        doc,
        ["หลักการ", "ข้อกำหนดในการพัฒนา"],
        [
            ("Evidence-bound", "AI ตอบจาก KPI, alert, policy document และ data source ที่ระบบส่งให้เท่านั้น"),
            ("Structured Answer", "ทุกคำตอบต้องมี Answer, Evidence, Why, Limitations, Next Action"),
            ("No Rights Decision", "ห้ามสรุปเชิงตัดสิทธิหรือชี้ว่าเด็กคนใดจะหลุดระบบแน่นอน"),
            ("Permission-aware", "retrieval และ response ต้องเคารพ role, masking และ audit log"),
            ("Traceable", "บันทึก prompt template, model version, retrieved evidence, timestamp"),
            ("Evaluation", "มีชุดคำถามทดสอบ unsupported claim, citation error, sensitive leakage"),
        ],
        [1.55, 4.95],
    )

    add_heading(doc, "9. Definition of Done", 1)
    for item in [
        "ผู้บริหารเข้าใจสถานการณ์สำคัญจากหน้าแรกโดยไม่ต้องอธิบายทุกกราฟ",
        "คลิกจาก KPI/Map ไปถึง Alert Detail ได้ภายใน 2-3 ขั้น",
        "Alert ทุกตัวมี severity, priority score, trigger, drivers, evidence, data quality และ rule/model version",
        "AI summary อ้างอิงข้อมูลจริงในระบบ ระบุข้อจำกัดชัดเจน และแสดง model/prompt version ที่ใช้",
        "KPI ทุกตัว drill-down ได้ถึง source, definition, period, denominator, aggregation level และ data quality flag",
        "ผู้ใช้สร้าง Case/Action จาก Alert และติดตาม owner/status/SLA ได้",
        "มี RBAC, masking, audit log และ data lineage ในระดับที่สาธิต governance ได้",
        "มี seed data และ demo script ที่รันซ้ำได้สำหรับ stakeholder review",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "10. Immediate Next Actions", 1)
    add_table(
        doc,
        ["ลำดับ", "งาน", "Owner", "Output"],
        [
            ("1", "Kickoff 90 นาทีเพื่อยืนยัน scope Phase 1 และ demo storyline", "Product Lead", "Approved MVP scope"),
            ("2", "สร้าง issue backlog ตาม epic และ acceptance criteria ในเอกสารนี้", "PM/Tech Lead", "Development backlog"),
            ("3", "เตรียม seed data และ mock API contract", "Data/Backend", "Sample dataset + OpenAPI"),
            ("4", "ทำ wireframe 5 หน้าหลักและ component inventory", "UX/UI", "Clickable prototype spec"),
            ("5", "ตั้ง project skeleton และ CI", "Engineering", "Running app shell"),
            ("6", "สร้าง dashboard + alert vertical slice แรก", "Full stack", "End-to-end demo slice"),
        ],
        [0.55, 3.05, 1.15, 1.75],
    )

    doc.core_properties.title = "EDU Sentinel System Development Work Order v1.1"
    doc.core_properties.subject = "Actionable development brief for starting EDU Sentinel prototype implementation"
    doc.core_properties.author = "Codex"
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
